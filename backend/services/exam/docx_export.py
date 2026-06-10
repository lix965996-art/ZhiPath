from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SECTION_TITLES = {
    "single_choice": "一、选择题",
    "multiple_choice": "二、多选题",
    "true_false": "三、判断题",
    "short_answer": "四、简答题",
}
FONT_NAME = "宋体"
BLACK = RGBColor(0, 0, 0)


def build_exam_docx(exam: dict[str, Any], include_answers: bool = True) -> bytes:
    document = Document()
    _setup_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    _add_run(title, str(exam.get("title") or "个性化测试卷"), size=18, bold=True)

    _add_exam_meta(document, exam)

    document.add_paragraph("")
    questions = exam.get("questions", []) or []
    question_number = 1

    for q_type, section_title in SECTION_TITLES.items():
        section_questions = [q for q in questions if q.get("type") == q_type]
        if not section_questions:
            continue
        _add_section_title(document, section_title)
        for question in section_questions:
            _add_question(document, question, question_number)
            question_number += 1

    if include_answers:
        document.add_page_break()
        _add_section_title(document, "答案与解析")
        for index, question in enumerate(questions, start=1):
            answer_text = _format_answer(question)
            paragraph = document.add_paragraph()
            _format_paragraph(paragraph, first_line=False)
            _add_run(paragraph, f"{index}. 答案：", bold=True)
            _add_run(paragraph, answer_text)
            explanation = str(question.get("explanation") or "").strip()
            if explanation:
                explanation_paragraph = document.add_paragraph()
                _format_paragraph(explanation_paragraph, first_line=False)
                _add_run(explanation_paragraph, f"解析：{explanation}")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Title"):
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(11)
        style.font.color.rgb = BLACK
        _set_style_east_asia_font(style, FONT_NAME)


def _set_style_east_asia_font(style: Any, font_name: str) -> None:
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def _set_run_font(run: Any, size: float = 11, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def _add_run(paragraph: Any, text: str, size: float = 11, bold: bool = False) -> Any:
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    return run


def _format_paragraph(paragraph: Any, first_line: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(22)


def _add_exam_meta(document: Document, exam: dict[str, Any]) -> None:
    meta_table = document.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = True
    _remove_table_borders(meta_table)

    values = [
        ("科目", str(exam.get("subject", "综合学习"))),
        ("主题", str(exam.get("topic", "个性化学习"))),
        ("时间", f"{exam.get('duration_minutes', 45)} 分钟"),
        ("满分", f"{exam.get('total_score', 100)} 分"),
    ]
    for index, (label, value) in enumerate(values):
        cell = meta_table.cell(0, index)
        _format_cell(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(paragraph, f"{label}：", bold=True)
        _add_run(paragraph, value)

    student_values = ["姓名：____________", "班级：____________", "日期：____________", ""]
    for index, value in enumerate(student_values):
        cell = meta_table.cell(1, index)
        _format_cell(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if value:
            _add_run(paragraph, value)


def _format_cell(cell: Any) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)


def _remove_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _add_section_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    _add_run(paragraph, title, size=13, bold=True)


def _add_question(document: Document, question: dict[str, Any], number: int) -> None:
    score = question.get("score", 0)
    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, first_line=False)
    _add_run(paragraph, f"{number}. ", bold=True)
    _add_run(paragraph, str(question.get("question", "")))
    _add_run(paragraph, f"（{score} 分）")

    options = question.get("options", []) or []
    for index, option in enumerate(options):
        option_paragraph = document.add_paragraph()
        _format_paragraph(option_paragraph, first_line=False)
        option_paragraph.paragraph_format.left_indent = Pt(18)
        _add_run(option_paragraph, f"{chr(65 + index)}. ")
        _add_run(option_paragraph, str(option))

    if question.get("type") == "short_answer":
        answer_paragraph = document.add_paragraph()
        _format_paragraph(answer_paragraph, first_line=False)
        _add_run(answer_paragraph, "答：")
        for _ in range(3):
            line = document.add_paragraph()
            _format_paragraph(line, first_line=False)
            _add_run(line, "____________________________________________________________")


def _format_answer(question: dict[str, Any]) -> str:
    answer = question.get("answer", "")
    if question.get("type") == "true_false":
        return "正确" if answer is True else "错误"
    if question.get("type") in {"single_choice", "multiple_choice"}:
        return _format_choice_answer(answer)
    if isinstance(answer, list):
        return "、".join(str(item) for item in answer)
    return str(answer)


def _format_choice_answer(answer: Any) -> str:
    if isinstance(answer, list):
        return "、".join(_choice_label(item) for item in answer)
    return _choice_label(answer)


def _choice_label(value: Any) -> str:
    if isinstance(value, int):
        return chr(65 + value) if 0 <= value < 26 else str(value)
    text = str(value).strip()
    if text.isdigit():
        index = int(text)
        return chr(65 + index) if 0 <= index < 26 else text
    return text
